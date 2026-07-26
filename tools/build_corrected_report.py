from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\joser\Documents\ExaToleraFallo\sistema-entradas")
OUT = ROOT / "output"
ASSETS = ROOT / "tmp" / "report_assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

DOCX_OUT = OUT / "Informe_Parte_V_Corregido.docx"

# Diseño resuelto:
# - preset: standard_business_brief
# - portada: editorial_cover
# - override académico: encabezados de tabla en azul para conservar
#   la identidad visual del informe original.

NAVY = "17365D"
BLUE = "1F4E78"
LIGHT_BLUE = "EAF2F8"
VERY_LIGHT = "F4F7FA"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "000000"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_paragraph(doc, text="", bold_prefix=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=NAVY)
    set_font(p.add_run(text), color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, VERY_LIGHT)
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Block"]
    for index, line in enumerate(dedent(text).strip().splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, name="Consolas", size=8.3, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=9.2, bold=True, color=WHITE)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), size=9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def make_diagram(path, title, nodes, arrows):
    width, height = 1700, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 34)
        node_font = ImageFont.truetype("arial.ttf", 25)
        small_font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = node_font = small_font = ImageFont.load_default()
    draw.text((width // 2, 35), title, font=title_font, fill=(23, 54, 93), anchor="ma")

    positions = {}
    for key, x, y, w, h, text, fill in nodes:
        positions[key] = (x, y, w, h)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline=(31, 78, 120), width=3)
        lines = text.split("\n")
        line_height = 31
        start_y = y + h / 2 - (len(lines) - 1) * line_height / 2
        for idx, line in enumerate(lines):
            draw.text((x + w / 2, start_y + idx * line_height), line, font=node_font, fill=(20, 35, 50), anchor="mm")

    for source, target, label in arrows:
        sx, sy, sw, sh = positions[source]
        tx, ty, tw, th = positions[target]
        start = (sx + sw, sy + sh / 2)
        end = (tx, ty + th / 2)
        draw.line((start, end), fill=(70, 90, 110), width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)], fill=(70, 90, 110))
        if label:
            draw.text(((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 - 18), label, font=small_font, fill=(70, 70, 70), anchor="mm")
    img.save(path, quality=95)


def build():
    overload_diagram = ASSETS / "sobrecarga_capas.png"
    db_diagram = ASSETS / "postgresql_ha.png"
    make_diagram(
        overload_diagram,
        "Defensa en capas frente a sobrecarga",
        [
            ("clients", 40, 285, 250, 130, "Clientes\nPico de tráfico", (255, 238, 238)),
            ("gateway", 380, 260, 300, 180, "API Gateway\nRate limit distribuido\nControl de admisión", (255, 248, 220)),
            ("pools", 790, 130, 320, 170, "Bulkheads\nColas acotadas\nCompra / consulta", (230, 244, 250)),
            ("pods", 1210, 130, 300, 170, "Gateway + Reservas\nHPA + límites", (230, 244, 250)),
            ("db", 1210, 430, 300, 150, "PostgreSQL\nPool limitado", (246, 232, 246)),
        ],
        [
            ("clients", "gateway", "HTTP"),
            ("gateway", "pools", "admitir o rechazar"),
            ("pools", "pods", "trabajo aceptado"),
        ],
    )
    make_diagram(
        db_diagram,
        "PostgreSQL de alta disponibilidad con escritor único",
        [
            ("app", 40, 285, 270, 130, "Reservas e\nInventario", (230, 244, 250)),
            ("proxy", 400, 250, 300, 200, "Proxy / Pool\nTimeouts\nCircuit Breaker", (255, 248, 220)),
            ("primary", 820, 100, 310, 170, "PostgreSQL primario\nEscrituras", (255, 235, 235)),
            ("replica", 820, 470, 310, 150, "PostgreSQL réplica\nWAL", (225, 245, 235)),
            ("dcs", 1260, 270, 330, 180, "Patroni / DCS\nLeader Election\nQuorum + fencing", (242, 232, 249)),
        ],
        [
            ("app", "proxy", "SQL"),
            ("proxy", "primary", "escrituras"),
            ("primary", "dcs", "lease"),
            ("replica", "dcs", "candidata"),
        ],
    )

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("Sistemas Distribuidos | Parte V - Análisis y Diseño"), size=8.5, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("SISTEMAS DISTRIBUIDOS"), size=12, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    set_font(p.add_run("SISTEMA DE RESERVAS DE ENTRADAS\nEN KUBERNETES"), size=24, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    set_font(p.add_run("Tolerancia a Fallos\nParte V - Análisis y Diseño"), size=14, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    set_font(p.add_run("INFORME TÉCNICO CORREGIDO"), size=13, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    set_font(p.add_run("AUTORES\nJosé Vanegas\nMiguel Vanegas"), size=10.5, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    set_font(p.add_run("Julio de 2026"), size=10, color=GRAY)

    doc.add_page_break()
    doc.add_heading("Resumen ejecutivo", level=1)
    add_paragraph(
        doc,
        "En la Parte II se catalogaron seis puntos de fallo. La Parte III seleccionó cuatro para implementar defensas prácticas: Inventario Fantasma, Pasarela Lenta, Correo Perdido y Condición de Carrera. Por tanto, los dos escenarios restantes analizados aquí son el Fallo 3, El Diluvio de Peticiones, y el Fallo 4, Base de Datos Intermitente."
    )
    add_paragraph(
        doc,
        "La sobrecarga se explica mediante teoría de colas, capacidad finita y backpressure. Su defensa de producción combina rate limiting distribuido, control de admisión, bulkheads, colas acotadas, autoscaling y observabilidad. La conectividad intermitente con PostgreSQL introduce incertidumbre sobre el resultado de las transacciones; requiere timeouts, retries limitados e idempotentes, Circuit Breaker y una arquitectura de alta disponibilidad con escritor único, quorum y fencing."
    )
    add_callout(
        doc,
        "Alcance",
        "Los dos fallos fueron provocados y observados durante la Parte II, pero no recibieron una defensa completa en la Parte III. Este informe distingue la evidencia observada de la solución propuesta para producción.",
    )
    doc.add_heading("Contenido", level=2)
    for item in (
        "Relación y trazabilidad con las Partes II y III",
        "Fallo 3: El Diluvio de Peticiones",
        "Fallo 4: Base de Datos Intermitente",
        "Relación con la investigación MLR de Leader Election",
        "Comparación, prioridades y conclusiones",
        "Referencias y anexos",
    ):
        add_number(doc, item)

    doc.add_page_break()
    doc.add_heading("1. Relación y trazabilidad con las Partes II y III", level=1)
    add_paragraph(
        doc,
        "La numeración utilizada en este informe conserva el catálogo original de la Parte II. Esto permite seguir cada escenario desde su mecanismo de inyección hasta la defensa implementada o la propuesta teórica."
    )
    add_table(
        doc,
        ["N.º", "Escenario", "Tratamiento", "Evidencia o salida"],
        [
            ("1", "Inventario Fantasma", "Implementado en Parte III", "Autorrecuperación y réplica disponible"),
            ("2", "Pasarela Lenta", "Implementado en Parte III", "Timeout, Circuit Breaker y compensación"),
            ("3", "Diluvio de Peticiones", "Analizado en Parte V", "k6; degradación sin errores HTTP"),
            ("4", "Base de Datos Intermitente", "Analizado en Parte V", "Service sin endpoints; DATABASE_ERROR"),
            ("5", "Correo Perdido", "Implementado en Parte III", "Retries, backoff y fallback"),
            ("6", "Condición de Carrera", "Implementado en Parte III", "UPDATE SQL atómico; una compra y un 409"),
        ],
        [600, 2450, 2500, 3810],
    )
    add_callout(
        doc,
        "Criterio de auditoría",
        "En las secciones siguientes, “resultado observado” se refiere a evidencia obtenida durante la inyección del fallo; “resultado esperado” se reserva para el diseño de producción aún no implementado.",
    )

    doc.add_heading("2. Fallo 3: El Diluvio de Peticiones", level=1)
    doc.add_heading("2.1 Descripción específica y evidencia observada", level=2)
    add_paragraph(
        doc,
        "El escenario ocurre cuando la tasa de solicitudes que llega al API Gateway supera la capacidad conjunta de Gateway, Reservas, Inventario y PostgreSQL. La prueba documentada en evidencias/fallo3-resultados-k6.txt utilizó 20 usuarios virtuales durante aproximadamente 30 segundos. Se atendieron 1433 solicitudes sin errores HTTP; sin embargo, la latencia máxima se aproximó a un segundo y el percentil 95 a medio segundo. El sistema mantuvo disponibilidad bajo esa carga, pero mostró degradación de rendimiento."
    )
    doc.add_heading("2.2 Explicación teórica", level=2)
    add_paragraph(
        doc,
        "Si lambda es la tasa de llegada, mu la tasa de servicio de una instancia y c el número de instancias, la utilización aproximada es rho = lambda/(c x mu). Cuando rho se acerca a 1, el tiempo de espera y la longitud de la cola crecen rápidamente; si lambda supera la capacidad de forma sostenida, la cola deja de vaciarse. El resultado puede ser agotamiento de threads, sockets, memoria y conexiones de base de datos."
    )
    add_paragraph(
        doc,
        "La acumulación genera fallos en cascada: el Gateway retiene solicitudes, Reservas conserva tareas pendientes y cada tarea puede ocupar conexiones HTTP y de PostgreSQL. CAP no es el modelo principal porque el problema no exige una partición de red ni una elección entre consistencia y disponibilidad de datos replicados. El fundamento principal es capacidad finita, teoría de colas y backpressure."
    )
    doc.add_heading("2.3 Causa raíz", level=2)
    for item in (
        "El API Gateway acepta tráfico sin cuota global ni cuota por cliente.",
        "Las colas y pools no tienen límites explícitos.",
        "El escalado es reactivo y puede llegar después del pico.",
        "Escalar solo el Gateway traslada la presión a Reservas o PostgreSQL.",
        "No existe separación de capacidad entre compras, consultas y endpoints administrativos.",
    ):
        add_bullet(doc, item)

    doc.add_heading("2.4 Solución propuesta para producción", level=2)
    add_paragraph(
        doc,
        "La primera defensa es un rate limiter distribuido en el Ingress/API Gateway. Un token bucket global y cuotas por cliente rechazan el exceso con HTTP 429 y Retry-After. Como Gateway tiene varias réplicas, el estado no debe mantenerse únicamente en memoria de cada pod: puede aplicarse en un proxy compartido, coordinarse mediante Redis o dividirse explícitamente la cuota entre instancias [3], [4]."
    )
    add_paragraph(
        doc,
        "Después se aplica control de admisión y bulkheads. Compras y consultas usan pools y colas acotados separados. Cuando una cola alcanza su capacidad, el sistema rechaza rápidamente con 503 en lugar de retener solicitudes indefinidamente. El HPA amplía Gateway y Reservas usando CPU y métricas de aplicación; para métricas personalizadas se requiere el adaptador correspondiente. Además, los pods deben declarar resource requests para que la utilización de CPU sea calculable [3]."
    )
    doc.add_picture(str(overload_diagram), width=Inches(6.35))
    doc.inline_shapes[-1]._inline.docPr.set(
        "descr",
        "Defensa en capas: clientes, API Gateway con rate limiting "
        "distribuido, bulkheads, HPA y PostgreSQL con pool limitado.",
    )
    p = doc.add_paragraph("Figura 1. Defensa en capas para contener la sobrecarga sin propagarla.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)

    doc.add_heading("2.5 Pseudocódigo de control de admisión", level=2)
    add_code(
        doc,
        """
        procesar_solicitud(request):
            clave = request.usuario_id o request.ip
            if not rate_limiter_distribuido.permitir_global():
                return 429, Retry-After: 2
            if not rate_limiter_distribuido.permitir_cliente(clave):
                return 429, Retry-After: 5

            pool = pool_compra if request.es_compra else pool_consulta
            if not pool.intentar_admitir():
                return 503, "CAPACIDAD_TEMPORALMENTE_AGOTADA"
            return ejecutar_con_timeout(pool, request, 3 segundos)
        """,
    )
    doc.add_heading("2.6 Resultado esperado y trade-offs", level=2)
    add_paragraph(
        doc,
        "El sistema puede rechazar una fracción del tráfico con 429 o 503, pero conserva tiempos de respuesta predecibles y evita una caída total. Se sacrifica disponibilidad para solicitudes que exceden la capacidad instantánea a cambio de proteger las operaciones admitidas. Los riesgos residuales son límites injustos, picos más rápidos que el escalado y dependencia del almacén compartido del rate limiter."
    )

    doc.add_heading("3. Fallo 4: Base de Datos Intermitente", level=1)
    doc.add_heading("3.1 Descripción específica y evidencia observada", level=2)
    add_paragraph(
        doc,
        "El fallo aparece cuando Reservas o Inventario pierden conectividad temporal y repetida con PostgreSQL durante una escritura. El script chaos/04-base-datos-intermitente.ps1 cambia temporalmente el selector de postgres-service: el pod permanece Running, pero el Service queda sin endpoints y los servicios dependientes devuelven DATABASE_ERROR. La evidencia demuestra que un proceso saludable no garantiza que la ruta de red, el Service o su backend sean alcanzables."
    )
    add_paragraph(
        doc,
        "Una desconexión puede ocurrir antes de enviar la transacción, durante su ejecución o después del COMMIT pero antes de que la aplicación reciba la confirmación. El último caso es especialmente peligroso: el cliente no sabe si la operación ocurrió y un retry ciego puede duplicar la reserva."
    )
    doc.add_heading("3.2 Explicación teórica y relación con CAP", level=2)
    add_paragraph(
        doc,
        "En un sistema asíncrono, un cliente no distingue de inmediato entre un servidor caído y una respuesta retrasada. Los timeouts convierten esa incertidumbre en una decisión operativa. Con una sola instancia de PostgreSQL, la pérdida de conectividad elimina disponibilidad; CAP se vuelve relevante al incorporar réplicas."
    )
    add_paragraph(
        doc,
        "Durante una partición, permitir escrituras en ambos lados puede generar primarios concurrentes y reservas contradictorias. Para un sistema de entradas debe priorizarse consistencia y tolerancia a particiones: solo el lado que conserva liderazgo válido puede aceptar escrituras; el otro rechaza temporalmente. Este comportamiento CP evita split-brain [2]."
    )

    doc.add_heading("3.3 Solución propuesta para producción", level=2)
    for item in (
        "PostgreSQL primario-réplica mediante streaming replication. La replicación síncrona reduce el RPO, pero aumenta latencia y puede reducir disponibilidad si no existe una réplica confirmando [5].",
        "Leader Election y fencing mediante Patroni u otro operador coordinado por un DCS. La promoción exige liderazgo válido y una réplica suficientemente actualizada [6].",
        "Proxy y pool de conexiones con connect_timeout, acquisition timeout y statement_timeout.",
        "Retries limitados solo para errores transitorios, con backoff exponencial y jitter.",
        "Idempotencia obligatoria para escrituras críticas y restricción UNIQUE en la base.",
        "Circuit Breaker y degradación controlada: las escrituras reciben 503 cuando no existe un primario seguro.",
    ):
        add_bullet(doc, item)
    doc.add_picture(str(db_diagram), width=Inches(6.35))
    doc.inline_shapes[-1]._inline.docPr.set(
        "descr",
        "PostgreSQL de alta disponibilidad: proxy, primario, réplica "
        "WAL y Patroni con DCS, liderazgo y fencing.",
    )
    p = doc.add_paragraph("Figura 2. Alta disponibilidad con escritor único, quorum y fencing.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading("3.4 Pseudocódigo corregido de compra idempotente", level=2)
    add_paragraph(
        doc,
        "La clave se reclama con un INSERT atómico. A diferencia de SELECT ... FOR UPDATE sobre una fila inexistente, esta operación impide que dos solicitudes nuevas avancen simultáneamente con la misma clave."
    )
    add_code(
        doc,
        """
        reservar_entrada(request, key):
            if circuit_breaker_db == OPEN:
                return 503, "DATABASE_TEMPORARILY_UNAVAILABLE"

            for intento in 1..3:
                try:
                    BEGIN
                    creada = INSERT INTO idempotency_keys(clave, estado)
                             VALUES (key, 'PROCESSING')
                             ON CONFLICT (clave) DO NOTHING
                             RETURNING clave

                    if no creada:
                        anterior = SELECT estado, resultado
                                   FROM idempotency_keys WHERE clave = key
                        ROLLBACK
                        if anterior.estado == 'COMPLETED':
                            return anterior.resultado
                        return 409, "REQUEST_ALREADY_IN_PROGRESS"

                    restante = UPDATE inventory
                               SET available_seats = available_seats - 1
                               WHERE event_id = request.event_id
                                 AND available_seats > 0
                               RETURNING available_seats
                    if no restante:
                        UPDATE idempotency_keys SET estado='COMPLETED',
                               resultado='409 SOLD_OUT' WHERE clave=key
                        COMMIT
                        return 409, "SOLD_OUT"

                    reserva = INSERT INTO reservations(...)
                    INSERT INTO outbox(tipo, payload)
                    UPDATE idempotency_keys SET estado='COMPLETED',
                           resultado=reserva WHERE clave=key
                    COMMIT
                    registrar_exito()
                    return 200, reserva
                except error_transitorio:
                    ROLLBACK
                    esperar(backoff(intento) + jitter)

            registrar_fallo()
            return 503, "DATABASE_TEMPORARILY_UNAVAILABLE"
        """,
    )
    add_callout(
        doc,
        "Nota de producción",
        "Los retries deben ejecutarse únicamente cuando la operación es idempotente. Una operación externa, como el cobro, debe usar su propia Idempotency-Key y coordinarse mediante outbox/saga; no debe repetirse ciegamente dentro de la transacción de inventario.",
    )

    doc.add_heading("3.5 Pseudocódigo conceptual de failover", level=2)
    add_code(
        doc,
        """
        controlador_ha():
            if el_líder no renueva su lease:
                if DCS no puede garantizar liderazgo seguro:
                    rechazar_escrituras()
                    return
                candidato = réplica_elegible_con_menor_WAL_lag()
                if candidato adquiere el lock de líder:
                    aplicar_fencing_al_primario_anterior()
                    promover(candidato)
                    actualizar_endpoint_de_escritura()
        """,
    )
    doc.add_heading("3.6 Resultado esperado y trade-offs", level=2)
    add_paragraph(
        doc,
        "Durante una desconexión breve, los retries limitados pueden completar la operación. Si la interrupción continúa, el Circuit Breaker evita agotar threads y conexiones. Ante la pérdida del primario, el sistema ejecuta failover y redirige escrituras al nuevo líder. El costo es una ventana de indisponibilidad, mayor complejidad y la necesidad de vigilar WAL lag, quorum, RTO, RPO y fencing."
    )
    add_callout(
        doc,
        "Decisión de consistencia",
        "Si no puede demostrarse quién posee el liderazgo válido, la respuesta correcta es rechazar temporalmente la compra; no es aceptable vender el último asiento en dos particiones.",
    )

    doc.add_heading("4. Relación con la investigación MLR de Leader Election", level=1)
    add_paragraph(
        doc,
        "La MLR previa del equipo analizó cómo Leader Election contribuye a coordinación, consistencia y alta disponibilidad, junto con sus costos de latencia y complejidad [7]. La conexión con PostgreSQL intermitente es directa: una arquitectura primaria-réplica necesita decidir qué instancia posee el derecho exclusivo de aceptar escrituras."
    )
    add_paragraph(
        doc,
        "Patroni utiliza un Distributed Configuration Store para coordinar el leader lock y detectar la pérdida del líder [6]. La elección por sí sola no es suficiente: el antiguo primario debe quedar aislado mediante fencing antes de que el nuevo escritor se considere seguro. Así se protege la unicidad del escritor y la consistencia del inventario."
    )

    doc.add_heading("5. Comparación y prioridades", level=1)
    add_table(
        doc,
        ["Fallo", "Fundamento", "Defensa prioritaria", "Respuesta", "Riesgo residual"],
        [
            ("3. Diluvio", "Colas, capacidad, backpressure", "Rate limit distribuido + bulkhead + HPA", "429/503 controlado", "Picos más rápidos que el escalado"),
            ("4. BD intermitente", "Incertidumbre, CAP e idempotencia", "HA + liderazgo + retry limitado", "503 o resultado idempotente", "Ventana de failover y posible pérdida asíncrona"),
        ],
        [1250, 1800, 2700, 1450, 2160],
    )
    doc.add_heading("Orden recomendado", level=2)
    for item in (
        "Añadir métricas, resource requests/limits, timeouts y dashboards.",
        "Implementar rate limiting distribuido y colas acotadas.",
        "Configurar HPA en Gateway y Reservas con métricas de recursos y aplicación.",
        "Incorporar claves de idempotencia y retries limitados en escrituras críticas.",
        "Migrar PostgreSQL a alta disponibilidad con failover, quorum y fencing.",
        "Ejecutar chaos testing periódico y medir RTO, RPO y SLO.",
    ):
        add_number(doc, item)

    doc.add_heading("6. Conclusiones", level=1)
    for item in (
        "El Diluvio de Peticiones es el Fallo 3 del catálogo y ocurre cuando la llegada supera la capacidad de procesamiento. El autoscaling ayuda, pero debe acompañarse de control de admisión y backpressure.",
        "La Base de Datos Intermitente es el Fallo 4 y puede dejar incierto el resultado de una transacción. Los retries sin idempotencia pueden duplicar operaciones.",
        "La idempotencia debe reclamarse de forma atómica; bloquear una fila inexistente no impide la carrera.",
        "Con replicación y una partición, el sistema prioriza consistencia: solo el escritor con liderazgo válido acepta compras.",
        "La solución de producción combina respuestas controladas, observabilidad, idempotencia, alta disponibilidad y pruebas periódicas de recuperación.",
    ):
        add_number(doc, item)

    doc.add_heading("Referencias", level=1)
    references = [
        "[1] Equipo del proyecto. Informe y evidencias del Sistema Distribuido de Reservas de Entradas. Partes II, III y IV, 2026.",
        "[2] Gilbert, S. y Lynch, N. Brewer's Conjecture and the Feasibility of Consistent, Available, Partition-Tolerant Web Services. ACM SIGACT News, 33(2), 51-59, 2002. DOI: 10.1145/564585.564601.",
        "[3] Kubernetes. Horizontal Pod Autoscaling. https://kubernetes.io/docs/concepts/workloads/autoscaling/horizontal-pod-autoscale/ (consulta: julio de 2026).",
        "[4] Envoy Proxy. HTTP Local Rate Limit Filter y Token Bucket. https://www.envoyproxy.io/docs/envoy/latest/configuration/http/http_filters/local_rate_limit_filter (consulta: julio de 2026).",
        "[5] PostgreSQL Global Development Group. High Availability, Load Balancing, and Replication. https://www.postgresql.org/docs/current/high-availability.html (consulta: julio de 2026).",
        "[6] Patroni. DCS Failsafe Mode y Leader Election. https://patroni.readthedocs.io/en/latest/dcs_failsafe_mode.html (consulta: julio de 2026).",
        "[7] Vanegas, J. y Vanegas, M. Leader Election en Sistemas Distribuidos: mecanismo de coordinación y alta disponibilidad ante fallas de nodos. Tarea MLR, 2026.",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_font(p.add_run(reference), size=9.5)

    doc.add_heading("Anexo A. HPA conceptual para Gateway", level=1)
    add_code(
        doc,
        """
        apiVersion: autoscaling/v2
        kind: HorizontalPodAutoscaler
        metadata:
          name: gateway-hpa
          namespace: tickets
        spec:
          scaleTargetRef:
            apiVersion: apps/v1
            kind: Deployment
            name: gateway
          minReplicas: 2
          maxReplicas: 10
          metrics:
            - type: Resource
              resource:
                name: cpu
                target:
                  type: Utilization
                  averageUtilization: 65
          behavior:
            scaleUp:
              stabilizationWindowSeconds: 0
              policies:
                - type: Percent
                  value: 100
                  periodSeconds: 60
            scaleDown:
              stabilizationWindowSeconds: 300
        """,
    )
    add_paragraph(
        doc,
        "Para utilizar CPU por porcentaje, el contenedor debe declarar resources.requests.cpu. Las métricas de RPS, longitud de cola o latencia p95 requieren custom.metrics.k8s.io, external.metrics.k8s.io o una integración equivalente [3]."
    )
    doc.add_heading("Anexo B. Política operativa de PostgreSQL", level=1)
    for item in (
        "connect_timeout corto y medible; statement_timeout según el SLO de compra.",
        "Máximo de tres intentos para errores transitorios: 200 ms, 400 ms y 800 ms más jitter.",
        "Idempotency-Key obligatoria en POST /reservations y restricción UNIQUE.",
        "Circuit Breaker con umbral, ventana de evaluación y transición half-open.",
        "Failover únicamente con liderazgo seguro y fencing.",
        "Alertas sobre WAL lag, cambios de líder, conexiones, errores, RTO y RPO.",
    ):
        add_bullet(doc, item)

    core = doc.core_properties
    core.title = "Sistema de Reservas de Entradas en Kubernetes - Parte V corregida"
    core.subject = "Tolerancia a fallos: análisis y diseño"
    core.author = "José Vanegas; Miguel Vanegas"
    core.keywords = "Kubernetes, tolerancia a fallos, rate limiting, PostgreSQL, leader election"
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build()
