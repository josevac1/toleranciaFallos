from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_corrected_report import (
    ASSETS,
    BLUE,
    DOCX_OUT,
    GRAY,
    NAVY,
    add_bullet,
    add_callout,
    add_code,
    add_number,
    add_page_number,
    add_paragraph,
    add_table,
    configure_styles,
    make_diagram,
    set_font,
)


def add_figure(doc, path, caption, alt_text):
    doc.add_picture(str(path), width=Inches(6.25))
    doc.inline_shapes[-1]._inline.docPr.set("descr", alt_text)
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.runs[0], size=9, italic=True, color=GRAY)


def build():
    notification_diagram = ASSETS / "correo_perdido_outbox.png"
    race_diagram = ASSETS / "condicion_carrera_holds.png"

    make_diagram(
        notification_diagram,
        "Entrega confiable de notificaciones mediante Outbox",
        [
            ("reservas", 40, 275, 260, 150, "Reservas\nTransacción", (230, 244, 250)),
            ("db", 390, 250, 300, 200, "PostgreSQL\nReserva + Outbox\nmismo COMMIT", (255, 248, 220)),
            ("relay", 790, 275, 260, 150, "Outbox Relay\nRetries + backoff", (230, 244, 250)),
            ("broker", 1140, 275, 230, 150, "Broker durable\nACK / DLQ", (242, 232, 249)),
            ("email", 1460, 275, 200, 150, "Notificaciones\nIdempotente", (225, 245, 235)),
        ],
        [
            ("reservas", "db", "COMMIT"),
            ("db", "relay", "pendientes"),
            ("relay", "broker", "publicar"),
            ("broker", "email", "consumir"),
        ],
    )
    make_diagram(
        race_diagram,
        "Reserva temporal del último asiento",
        [
            ("clients", 40, 275, 250, 150, "Clientes A y B\nsolicitud simultánea", (255, 238, 238)),
            ("api", 380, 275, 260, 150, "Reservas\nIdempotency-Key", (230, 244, 250)),
            ("db", 750, 220, 360, 260, "PostgreSQL\nUPDATE condicional\nHOLD con expiración\nUNIQUE + transacción", (255, 248, 220)),
            ("winner", 1220, 120, 300, 150, "Ganador\nHOLD / CONFIRMED", (225, 245, 235)),
            ("loser", 1220, 480, 300, 150, "Segundo cliente\n409 SOLD_OUT", (255, 235, 235)),
        ],
        [
            ("clients", "api", "concurrencia"),
            ("api", "db", "transacción"),
            ("db", "winner", "1 fila"),
            ("db", "loser", "0 filas"),
        ],
    )

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(
        header.add_run("Sistemas Distribuidos | Parte V - Análisis y Diseño"),
        size=8.5,
        color=GRAY,
    )
    add_page_number(section.footer.paragraphs[0])

    for _ in range(5):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.add_run("SISTEMAS DISTRIBUIDOS"), size=12, bold=True, color=BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    set_font(
        paragraph.add_run("SISTEMA DE RESERVAS DE ENTRADAS\nEN KUBERNETES"),
        size=24,
        bold=True,
        color=NAVY,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    set_font(
        paragraph.add_run("Tolerancia a Fallos\nParte V - Análisis y Diseño"),
        size=14,
        color=GRAY,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(38)
    set_font(paragraph.add_run("INFORME TÉCNICO FINAL"), size=13, bold=True, color=BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(22)
    set_font(paragraph.add_run("AUTORES\nJosé Vanegas\nMiguel Vanegas"), size=10.5, color=NAVY)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    set_font(paragraph.add_run("Julio de 2026"), size=10, color=GRAY)

    doc.add_page_break()
    doc.add_heading("Resumen ejecutivo", level=1)
    add_paragraph(
        doc,
        "Este informe analiza dos fallos del sistema de reservas: el Fallo 5, El Correo Perdido, y el Fallo 6, Condición de Carrera. Para cada escenario se explica su causa teórica, se propone una solución de nivel producción y se presenta pseudocódigo y un diagrama.",
    )
    add_paragraph(
        doc,
        "El Correo Perdido es un fallo parcial de una operación secundaria. La solución de producción propuesta combina Transactional Outbox, mensajería durable, entrega al menos una vez, consumidores idempotentes, retries con backoff y Dead Letter Queue. La Condición de Carrera es un problema de concurrencia sobre estado compartido; se resuelve mediante una actualización condicional atómica o reservas temporales con expiración, restricciones de base de datos e idempotencia.",
    )
    add_callout(
        doc,
        "Alcance",
        "El análisis se limita a los servicios, datos y decisiones técnicas involucrados en estos dos fallos. Las propuestas distinguen las salvaguardas básicas del prototipo de una arquitectura completa de producción.",
    )
    doc.add_heading("Contenido", level=2)
    for item in (
        "Alcance y criterios de análisis",
        "Fallo 5: El Correo Perdido",
        "Fallo 6: Condición de Carrera",
        "Relación con fundamentos distribuidos y la MLR",
        "Comparación, prioridades y conclusiones",
        "Referencias",
    ):
        add_number(doc, item)

    doc.add_page_break()
    doc.add_heading("1. Alcance y criterios de análisis", level=1)
    add_paragraph(
        doc,
        "Los dos escenarios se evalúan con cuatro criterios: causa específica, impacto sobre el usuario, garantía que debe preservar el sistema y solución de producción. No se presentan como descripciones genéricas de tolerancia a fallos.",
    )
    add_table(
        doc,
        ["N.º", "Escenario", "Problema central", "Garantía requerida"],
        [
            ("5", "Correo Perdido", "Fallo parcial en una operación secundaria", "No perder la notificación ni revertir la compra"),
            ("6", "Condición de Carrera", "Concurrencia sobre el último asiento", "Una sola confirmación por asiento"),
        ],
        [700, 2150, 3050, 3460],
    )
    add_callout(
        doc,
        "Criterio",
        "Cada solución debe preservar una garantía explícita y mostrar sus costos, límites y riesgos residuales.",
    )

    doc.add_heading("2. Fallo 5: El Correo Perdido", level=1)
    doc.add_heading("2.1 Descripción específica", level=2)
    add_paragraph(
        doc,
        "El usuario completa el pago y obtiene su entrada, pero el Servicio de Notificaciones está inactivo o no confirma la entrega. El fallo no debe revertir una compra válida, aunque tampoco puede ignorarse: el usuario necesita recibir su comprobante y el equipo debe poder detectar y reprocesar los mensajes pendientes.",
    )
    doc.add_heading("2.2 Explicación teórica", level=2)
    add_paragraph(
        doc,
        "El escenario es un fallo parcial entre dos operaciones: confirmar la reserva en PostgreSQL y solicitar el correo. Si se guarda la reserva y luego falla la llamada HTTP, el estado principal queda correcto pero la notificación se pierde. Si se intenta enviar antes del COMMIT, puede enviarse un correo de una reserva que finalmente sea revertida. Una transacción ACID local no puede confirmar atómicamente una escritura SQL y una llamada remota.",
    )
    add_paragraph(
        doc,
        "CAP no es el modelo principal de este fallo; el problema central es la coordinación entre recursos con límites transaccionales diferentes. Una garantía práctica es entrega al menos una vez: el mensaje puede repetirse, por lo que el consumidor debe ser idempotente. La entrega exactamente una vez de extremo a extremo no puede suponerse solo porque el broker marque un mensaje como procesado.",
    )
    doc.add_heading("2.3 Solución propuesta para producción", level=2)
    for item in (
        "Transactional Outbox: la reserva y el evento NotificationRequested se escriben dentro del mismo COMMIT.",
        "Outbox Relay: un proceso publica eventos pendientes en un broker durable y marca cada fila como enviada solo después del ACK.",
        "Retries con backoff exponencial y jitter para errores transitorios; los errores permanentes terminan en una Dead Letter Queue.",
        "Consumidor idempotente mediante una clave única message_id para evitar correos duplicados.",
        "Estado visible CONFIRMED_NOTIFICATION_PENDING, métricas de antigüedad del outbox, alertas y herramienta de reproceso.",
    ):
        add_bullet(doc, item)
    add_figure(
        doc,
        notification_diagram,
        "Figura 1. Outbox transaccional y entrega durable de notificaciones.",
        "Flujo de Reservas a PostgreSQL y Outbox, relay, broker durable y Servicio de Notificaciones idempotente.",
    )
    doc.add_heading("2.4 Pseudocódigo", level=2)
    add_code(
        doc,
        """
        confirmar_reserva(request):
            BEGIN
            reserva = INSERT INTO reservations(..., status='CONFIRMED')
            INSERT INTO outbox(
                message_id, tipo, payload, estado
            ) VALUES (
                uuid(), 'NotificationRequested', reserva, 'PENDING'
            )
            COMMIT
            return 200, reserva

        relay_outbox():
            for evento in SELECT ... FROM outbox
                          WHERE estado='PENDING'
                          FOR UPDATE SKIP LOCKED:
                if broker.publicar(evento, persistent=true):
                    UPDATE outbox SET estado='PUBLISHED'
                else:
                    programar_retry_con_backoff_o_DLQ(evento)

        consumir_notificacion(evento):
            if message_id ya existe en processed_messages:
                ACK
                return
            enviar_correo(evento.payload)
            INSERT INTO processed_messages(message_id)
            ACK
        """,
    )
    doc.add_heading("2.5 Resultado esperado y trade-offs", level=2)
    add_paragraph(
        doc,
        "La compra responde CONFIRMED aunque Notificaciones esté caída. El evento permanece durable y se entrega cuando el servicio se recupera. Se acepta la posibilidad de duplicados controlados a cambio de no perder mensajes. Los costos son mayor complejidad operativa, almacenamiento del outbox, mantenimiento del relay y monitoreo de la DLQ.",
    )

    doc.add_heading("3. Fallo 6: Condición de Carrera", level=1)
    doc.add_heading("3.1 Descripción específica", level=2)
    add_paragraph(
        doc,
        "Dos usuarios consultan simultáneamente el último asiento. Si ambos leen available_seats = 1 y luego descuentan en operaciones separadas, los dos pueden recibir confirmación. El error no es la caída de un nodo, sino un interleaving inválido entre operaciones concurrentes.",
    )
    doc.add_heading("3.2 Explicación teórica", level=2)
    add_paragraph(
        doc,
        "La secuencia comprobar-luego-actuar es vulnerable a una actualización perdida o a una violación del invariante available_seats >= 0. El aislamiento READ COMMITTED no convierte automáticamente dos sentencias separadas en una decisión única. La corrección exige que la validación y el descuento formen una operación indivisible o que la fila se bloquee explícitamente dentro de una transacción.",
    )
    add_paragraph(
        doc,
        "CAP tampoco es el fundamento principal mientras todas las escrituras se resuelvan en un único PostgreSQL. El modelo relevante es control de concurrencia y serialización. Si el inventario se replicara entre regiones con escrituras concurrentes, entonces sí aparecería una decisión entre disponibilidad durante una partición y consistencia del último asiento.",
    )
    doc.add_heading("3.3 Solución propuesta para producción", level=2)
    for item in (
        "Actualización condicional atómica: UPDATE inventory SET available_seats = available_seats - 1 WHERE event_id = ? AND available_seats > 0 RETURNING available_seats.",
        "Restricción CHECK available_seats >= 0 como última defensa de integridad.",
        "Idempotency-Key única por intento de compra para que un retry no descuente otro asiento.",
        "Holds con expiración: el asiento pasa temporalmente a HELD mientras se procesa el pago y vuelve a AVAILABLE si vence el TTL.",
        "Índice o restricción UNIQUE que impida dos holds activos sobre el mismo asiento.",
        "Transacción corta; no mantener un bloqueo de base de datos mientras se espera una llamada externa de pago.",
    ):
        add_bullet(doc, item)
    add_figure(
        doc,
        race_diagram,
        "Figura 2. Resolución determinista de dos compras sobre el último asiento.",
        "Dos clientes compiten; una transacción crea el hold y la otra obtiene 409 SOLD_OUT.",
    )
    doc.add_heading("3.4 Pseudocódigo", level=2)
    add_code(
        doc,
        """
        crear_hold(request, idempotency_key):
            BEGIN
            if existe_resultado(idempotency_key):
                COMMIT
                return resultado_previo

            asiento = UPDATE inventory
                      SET available_seats = available_seats - 1
                      WHERE event_id = request.event_id
                        AND available_seats > 0
                      RETURNING available_seats
            if no asiento:
                guardar_resultado(idempotency_key, 409, 'SOLD_OUT')
                COMMIT
                return 409, 'SOLD_OUT'

            hold = INSERT INTO holds(
                event_id, user_id, status, expires_at
            ) VALUES (..., 'HELD', now() + 5 minutos)
            guardar_resultado(idempotency_key, 201, hold)
            COMMIT
            return 201, hold

        confirmar_hold(hold_id):
            UPDATE holds
            SET status='CONFIRMED'
            WHERE id=hold_id AND status='HELD'
              AND expires_at > now()

        expirar_holds():
            liberar inventario de holds vencidos
            usando una transición idempotente
        """,
    )
    doc.add_heading("3.5 Resultado esperado y trade-offs", level=2)
    add_paragraph(
        doc,
        "Solo una solicitud obtiene el último asiento; la otra recibe 409 SOLD_OUT. Los holds mejoran la experiencia durante el pago, pero introducen expiraciones, limpieza periódica y capacidad temporalmente retenida. El bloqueo pesimista es más simple para transacciones cortas, mientras que los holds evitan conservar bloqueos SQL durante dependencias lentas.",
    )

    doc.add_heading("4. Fundamentos y relación con la MLR", level=1)
    add_paragraph(
        doc,
        "La investigación MLR del equipo sobre Leader Election es útil para distinguir niveles del problema. Elegir un líder evita múltiples coordinadores o escritores en una arquitectura replicada, pero no sustituye el control transaccional dentro del líder. Un único primario todavía puede procesar dos solicitudes concurrentes de forma incorrecta si la aplicación usa comprobar-luego-actuar.",
    )
    add_paragraph(
        doc,
        "En El Correo Perdido, Leader Election podría asegurar que solo una instancia del relay procese una partición del outbox, aunque la idempotencia sigue siendo obligatoria ante reintentos. En la Condición de Carrera, la elección de líder protege la unicidad del escritor entre nodos; la sentencia atómica, el bloqueo o el nivel de aislamiento protegen el invariante entre transacciones.",
    )

    doc.add_heading("5. Comparación y prioridades", level=1)
    add_table(
        doc,
        ["Fallo", "Fundamento", "Defensa principal", "Respuesta al usuario", "Riesgo residual"],
        [
            ("5. Correo Perdido", "Fallo parcial y entrega de mensajes", "Outbox + broker + consumidor idempotente", "Compra confirmada; correo pendiente", "Duplicados, atraso o DLQ"),
            ("6. Condición de Carrera", "Concurrencia y serialización", "UPDATE atómico + hold + idempotencia", "201 para uno; 409 para el otro", "Holds vencidos o contención"),
        ],
        [1400, 1800, 2600, 1900, 1660],
    )
    doc.add_heading("6. Conclusiones", level=1)
    for item in (
        "El Correo Perdido es el Fallo 5 y requiere desacoplar la confirmación de compra de la entrega del correo.",
        "La Condición de Carrera es el Fallo 6 y exige una decisión atómica sobre el inventario compartido.",
        "CAP y Leader Election aportan contexto cuando existen particiones o escritores replicados, pero no reemplazan idempotencia, mensajería durable ni control de concurrencia.",
        "Las soluciones propuestas preservan la operación principal y hacen visibles los estados pendientes, errores y riesgos residuales.",
    ):
        add_number(doc, item)

    doc.add_heading("Referencias", level=1)
    references = (
        "[1] Equipo del proyecto. Sistema Distribuido de Reservas de Entradas. Código, manifiestos y evidencias técnicas, 2026.",
        "[2] Gilbert, S. y Lynch, N. Brewer's Conjecture and the Feasibility of Consistent, Available, Partition-Tolerant Web Services. ACM SIGACT News, 2002. DOI: 10.1145/564585.564601.",
        "[3] PostgreSQL Global Development Group. Transaction Isolation. https://www.postgresql.org/docs/current/transaction-iso.html (consulta: julio de 2026).",
        "[4] PostgreSQL Global Development Group. Explicit Locking. https://www.postgresql.org/docs/current/explicit-locking.html (consulta: julio de 2026).",
        "[5] Richardson, C. Transactional Outbox Pattern. https://microservices.io/patterns/data/transactional-outbox.html (consulta: julio de 2026).",
        "[6] Cloud Native Computing Foundation. CloudEvents Specification. https://cloudevents.io/ (consulta: julio de 2026).",
        "[7] Vanegas, J. y Vanegas, M. Leader Election en Sistemas Distribuidos. Tarea MLR, 2026.",
    )
    for reference in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        set_font(paragraph.add_run(reference), size=9.5)

    properties = doc.core_properties
    properties.title = "Sistema de Reservas de Entradas - Parte V final"
    properties.subject = "Parte V - Análisis y Diseño"
    properties.author = "José Vanegas; Miguel Vanegas"
    properties.keywords = "Kubernetes, transactional outbox, concurrencia, idempotencia"
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build()
