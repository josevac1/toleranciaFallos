from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_corrected_report import (
    BLUE,
    GRAY,
    NAVY,
    add_bullet,
    add_callout,
    add_code,
    add_number,
    add_page_number,
    add_paragraph,
    add_table,
    configure_styles,
    set_font,
)


ROOT = Path(r"C:\Users\joser\Documents\ExaToleraFallo\sistema-entradas")
OUTPUT = ROOT / "output"
DOCX_OUTPUT = OUTPUT / "Informe_Parte_II_Corregido.docx"
ARCHITECTURE_IMAGE = ROOT / "evidencias" / "diagrama.png"


def add_evidence_image(doc, path, caption, alt_text):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.25))
    doc.inline_shapes[-1]._inline.docPr.set("descr", alt_text)
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.runs[0], size=9, italic=True, color=GRAY)


def build():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(
        header.add_run("Sistemas Distribuidos | Parte II - Seis Puntos de Fallo"),
        size=8.5,
        color=GRAY,
    )
    add_page_number(section.footer.paragraphs[0])

    for _ in range(5):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.add_run("SISTEMAS DISTRIBUIDOS"), size=12, bold=True, color=BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    set_font(
        paragraph.add_run("SISTEMA DE RESERVAS DE ENTRADAS\nEN KUBERNETES"),
        size=24,
        bold=True,
        color=NAVY,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    set_font(
        paragraph.add_run("Parte II - Los 6 Puntos de Fallo\nChaos Scenarios"),
        size=14,
        color=GRAY,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(38)
    set_font(paragraph.add_run("INFORME TÉCNICO CORREGIDO"), size=13, bold=True, color=BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(22)
    set_font(paragraph.add_run("AUTORES\nJosé Vanegas\nMiguel Vanegas"), size=10.5, color=NAVY)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    set_font(paragraph.add_run("Julio de 2026"), size=10, color=GRAY)

    doc.add_page_break()
    doc.add_heading("1. Objetivo y alcance", level=1)
    add_paragraph(
        doc,
        "Comprender el catálogo de los seis escenarios de fallo del sistema de reservas y establecer, antes de su experimentación, qué mecanismo técnico permitiría provocar cada uno de forma controlada sobre el clúster Kubernetes.",
    )
    doc.add_heading("2. Arquitectura utilizada", level=1)
    add_paragraph(
        doc,
        "El flujo principal es Cliente -> API Gateway -> Reservas -> Inventario -> Pagos -> Notificaciones, con PostgreSQL como persistencia. Gateway, Reservas e Inventario tienen dos réplicas distribuidas entre los nodos del clúster tickets-cluster.",
    )
    add_evidence_image(
        doc,
        ARCHITECTURE_IMAGE,
        "Figura 1. Arquitectura desplegada para los experimentos de caos.",
        "Arquitectura del sistema de reservas distribuida entre dos nodos Kubernetes.",
    )

    doc.add_heading("3. Catálogo y mecanismo de inyección", level=1)
    add_table(
        doc,
        ["N.º", "Punto de fallo", "Tipo", "Mecanismo controlado"],
        [
            ("1", "Inventario Fantasma", "Disponibilidad", "Eliminar un pod de Inventario"),
            ("2", "Pasarela Lenta", "Latencia", "Pagos responde después de 20 s"),
            ("3", "Diluvio de Peticiones", "Sobrecarga", "Job de k6 con pico de usuarios"),
            ("4", "Base de Datos Intermitente", "Conectividad", "Service de PostgreSQL sin endpoints"),
            ("5", "Correo Perdido", "Fallo no crítico", "Notificaciones con cero réplicas"),
            ("6", "Condición de Carrera", "Consistencia", "Dos compras simultáneas"),
        ],
        [600, 2150, 1600, 5010],
    )

    doc.add_heading("4. Criterio de aplicación", level=1)
    add_paragraph(
        doc,
        "El mecanismo indicado en la tabla define únicamente cómo generar cada anomalía de manera controlada. Esta clasificación no desarrolla patrones de defensa ni soluciones de producción; su propósito es dejar explícita la relación entre el fallo y su método de inyección.",
    )

    doc.add_heading("5. Conclusiones", level=1)
    for item in (
        "Los seis escenarios quedaron asociados a un mecanismo de inyección reproducible.",
        "El catálogo cubre disponibilidad, latencia, sobrecarga, conectividad, fallo no crítico y consistencia.",
        "La causa de cada anomalía y el método utilizado para provocarla quedan diferenciados de forma explícita.",
    ):
        add_number(doc, item)

    properties = doc.core_properties
    properties.title = "Sistema de Reservas de Entradas - Parte II corregida"
    properties.subject = "Parte II - Los 6 Puntos de Fallo"
    properties.author = "José Vanegas; Miguel Vanegas"
    properties.keywords = "Kubernetes, chaos engineering, seis fallos, tolerancia a fallos"
    doc.save(DOCX_OUTPUT)
    print(DOCX_OUTPUT)


if __name__ == "__main__":
    build()
