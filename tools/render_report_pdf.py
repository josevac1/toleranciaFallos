import argparse
from html import escape
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph as PdfParagraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"C:\Users\joser\Documents\ExaToleraFallo\sistema-entradas")
DOCX_INPUT = ROOT / "output" / "Informe_Parte_V_Corregido.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "Informe_Parte_V_Corregido.pdf"

NAVY = colors.HexColor("#17365D")
BLUE = colors.HexColor("#1F4E78")
LIGHT_BLUE = colors.HexColor("#EAF2F8")
VERY_LIGHT = colors.HexColor("#F4F7FA")
GRAY = colors.HexColor("#666666")


def iter_blocks(document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def paragraph_has_page_break(paragraph):
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def paragraph_image(paragraph, document):
    blips = paragraph._p.xpath(".//a:blip")
    if not blips:
        return None
    relationship_id = blips[0].get(qn("r:embed"))
    if not relationship_id:
        return None
    part = document.part.related_parts[relationship_id]
    stream = BytesIO(part.blob)
    image = Image(stream)
    max_width = 6.3 * inch
    max_height = 4.3 * inch
    scale = min(max_width / image.imageWidth, max_height / image.imageHeight, 1)
    image.drawWidth = image.imageWidth * scale
    image.drawHeight = image.imageHeight * scale
    image.hAlign = "CENTER"
    return image


def build_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "BodyAcademic",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=13.2,
            spaceAfter=6,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            "HeadingAcademic1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            textColor=BLUE,
            spaceBefore=16,
            spaceAfter=8,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            "HeadingAcademic2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=BLUE,
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            "HeadingAcademic3",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            "BulletAcademic",
            parent=styles["BodyAcademic"],
            leftIndent=0.5 * inch,
            firstLineIndent=-0.25 * inch,
            bulletIndent=0.25 * inch,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "CodeAcademic",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7.4,
            leading=9.2,
            leftIndent=6,
            rightIndent=6,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            "CaptionAcademic",
            parent=styles["BodyAcademic"],
            fontName="Helvetica-Oblique",
            fontSize=8.7,
            leading=10.5,
            alignment=TA_CENTER,
            textColor=GRAY,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            "ReferenceAcademic",
            parent=styles["BodyAcademic"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=10.5,
            leftIndent=0.2 * inch,
            firstLineIndent=-0.2 * inch,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            "CoverKicker",
            parent=styles["BodyAcademic"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=BLUE,
            alignment=TA_CENTER,
            spaceAfter=20,
        )
    )
    styles.add(
        ParagraphStyle(
            "CoverTitle",
            parent=styles["BodyAcademic"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=27,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            "CoverSubtitle",
            parent=styles["BodyAcademic"],
            fontName="Helvetica",
            fontSize=14,
            leading=18,
            textColor=GRAY,
            alignment=TA_CENTER,
            spaceAfter=28,
        )
    )
    return styles


def table_to_flowable(docx_table, styles):
    rows = []
    one_cell = len(docx_table.rows) == 1 and len(docx_table.columns) == 1
    is_code = one_cell and any(
        paragraph.style and paragraph.style.name == "Code Block"
        for paragraph in docx_table.cell(0, 0).paragraphs
    )
    is_callout = one_cell and not is_code
    for row_index, row in enumerate(docx_table.rows):
        values = []
        for cell in row.cells:
            text = "\n".join(p.text for p in cell.paragraphs).strip()
            if is_code:
                markup = escape(text).replace("\n", "<br/>")
                values.append(PdfParagraph(markup, styles["CodeAcademic"]))
            else:
                markup = escape(text).replace("\n", "<br/>")
                style = ParagraphStyle(
                    f"Cell{row_index}",
                    parent=styles["BodyAcademic"],
                    fontName=(
                        "Helvetica"
                        if is_callout
                        else ("Helvetica-Bold" if row_index == 0 else "Helvetica")
                    ),
                    fontSize=8.2 if row_index == 0 else 8.0,
                    leading=10,
                    textColor=(
                        NAVY
                        if is_callout
                        else (colors.white if row_index == 0 else colors.black)
                    ),
                    alignment=TA_CENTER if row_index == 0 else TA_LEFT,
                    spaceAfter=0,
                )
                values.append(PdfParagraph(markup, style))
        rows.append(values)

    column_count = max(1, len(rows[0]))
    widths = [6.5 * inch / column_count] * column_count
    flowable = Table(
        rows,
        colWidths=widths,
        repeatRows=0 if one_cell else 1,
    )
    if is_code:
        flowable.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), VERY_LIGHT),
                    ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#D7DEE8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            )
        )
    elif is_callout:
        flowable.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BLUE),
                    ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D8E8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            )
        )
    else:
        flowable.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), BLUE),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB7C4")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
    return flowable


def header_footer(canvas, document, label):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(GRAY)
    canvas.drawRightString(
        letter[0] - inch,
        letter[1] - 0.55 * inch,
        f"Sistemas Distribuidos | {label}",
    )
    canvas.drawRightString(letter[0] - inch, 0.52 * inch, f"Página {document.page}")
    canvas.restoreState()


def first_page_footer(canvas, document):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(GRAY)
    canvas.drawRightString(letter[0] - inch, 0.52 * inch, f"Página {document.page}")
    canvas.restoreState()


def build_pdf(docx_input=DOCX_INPUT, pdf_output=PDF_OUTPUT):
    docx_input = Path(docx_input)
    pdf_output = Path(pdf_output)
    pdf_output.parent.mkdir(parents=True, exist_ok=True)
    document = Document(docx_input)
    header_label = (
        document.core_properties.subject
        or document.core_properties.title
        or "Informe técnico"
    )
    styles = build_styles()
    story = []
    cover_nonempty_index = 0
    on_cover = True
    list_number = 0

    for block in iter_blocks(document):
        if isinstance(block, DocxTable):
            story.append(table_to_flowable(block, styles))
            story.append(Spacer(1, 7))
            continue

        if paragraph_has_page_break(block):
            story.append(PageBreak())
            on_cover = False
            list_number = 0
            continue

        image = paragraph_image(block, document)
        if image is not None:
            story.append(image)
            story.append(Spacer(1, 4))
            continue

        text = block.text.strip()
        if not text:
            if not story:
                story.append(Spacer(1, 20))
            continue

        style_name = block.style.name if block.style else "Normal"
        markup = escape(text).replace("\n", "<br/>")
        if style_name.startswith("Heading "):
            list_number = 0

        if on_cover:
            cover_nonempty_index += 1
            if cover_nonempty_index == 1:
                style = styles["CoverKicker"]
            elif cover_nonempty_index == 2:
                style = styles["CoverTitle"]
            elif cover_nonempty_index == 3:
                style = styles["CoverSubtitle"]
            else:
                style = styles["BodyAcademic"]
                style = ParagraphStyle(
                    f"CoverBody{cover_nonempty_index}",
                    parent=style,
                    alignment=TA_CENTER,
                    textColor=NAVY if cover_nonempty_index < 6 else GRAY,
                    spaceBefore=12,
                )
        elif style_name == "Heading 1":
            style = styles["HeadingAcademic1"]
        elif style_name == "Heading 2":
            style = styles["HeadingAcademic2"]
        elif style_name == "Heading 3":
            style = styles["HeadingAcademic3"]
        elif style_name == "List Bullet":
            story.append(
                PdfParagraph(markup, styles["BulletAcademic"], bulletText="•")
            )
            continue
        elif style_name == "List Number":
            list_number += 1
            style = styles["BulletAcademic"]
            markup = f"<b>{list_number}.</b> {markup}"
        elif text.startswith("Figura "):
            style = styles["CaptionAcademic"]
        elif len(text) > 2 and text[0] == "[" and text[1].isdigit():
            style = styles["ReferenceAcademic"]
        else:
            style = styles["BodyAcademic"]

        story.append(PdfParagraph(markup, style))

    pdf = SimpleDocTemplate(
        str(pdf_output),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.82 * inch,
        bottomMargin=0.78 * inch,
        title="Sistema de Reservas de Entradas en Kubernetes - Parte V corregida",
        author="José Vanegas; Miguel Vanegas",
    )
    pdf.build(
        story,
        onFirstPage=first_page_footer,
        onLaterPages=lambda canvas, doc: header_footer(
            canvas,
            doc,
            header_label,
        ),
    )
    print(pdf_output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=DOCX_INPUT)
    parser.add_argument("--output", type=Path, default=PDF_OUTPUT)
    arguments = parser.parse_args()
    build_pdf(arguments.input, arguments.output)
